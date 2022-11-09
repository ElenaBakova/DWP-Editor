import json
import os
import re
import shutil
import sys

from docx import Document

from doc.Section import Section
from doc.WorkProgram import WorkProgram

dir_path = os.path.dirname(os.path.realpath(__file__)) + "\\..\\"


def read_docx(filename):
    doc = Document(filename)
    content = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text == '':
            continue
        content.append(text)

    return content


def parse_tables(file):
    doc = Document(file)
    tables_content = []

    for table in doc.tables:
        curr_table = []
        for row in table.rows:
            values = []
            for cell in row.cells:
                for para in cell.paragraphs:
                    values.append(para.text.strip())
            curr_table.append(values)
        tables_content.append(curr_table)

    return tables_content


def parse(doc):
    wp = WorkProgram()
    res = wp.content
    flag = False
    prev = '9.9.9.9.9'

    r = "Титульная страница"
    res[r] = Section()

    for par in doc:

        if par == '' or par == ' ':
            continue

        __match = re.match(r'Раздел \d\.', par)  # находим заголовки разделов

        if __match is not None:
            flag = False
            r = __match.group(0)
            res[r] = Section()
            res[r].heading = par[(__match.end() + 1):]
            continue

        _match = re.match(r'\n*\d\.\d\.\d?', par)  # находим заголовки подразделов

        if _match is not None:
            m = _match.group(0)

            res[m] = Section()
            res[m].heading = par[(_match.end() + 1):]

            if not flag:
                res[r].links.append(m)
                prev = m
            elif len(prev) < len(m):
                res[prev].links.append(m)
            else:
                res[r].links.append(m)
                prev = m
            flag = True
        elif flag:
            res[m].content = res[m].content + par + "\n"
        else:
            res[r].content = res[r].content + par + "\n"

    return res


'''''
def access(doc, heading):
    cont = doc.get(heading)
    assert cont is not None
    return cont
'''''


def feed_content(file):
    doc = read_docx(file)
    parsing = parse(doc)

    data = dict()
    for x in parsing:
        data[x] = {
            "title": parsing[x].heading,
            "text": parsing[x].content
        }
        '''''
    tables = parse_tables(file)
    data['table'] = list()
    for x in tables:
        for i in x:
            data['table'].append(i)
'''''
    return data


def feed_structure(file):
    doc = read_docx(file)
    parsing = parse(doc)
    data = dict()
    for x in parsing:
        data[x] = parsing[x].links
    return data


def feed(file):
    data = feed_structure(file)

    data1 = feed_content(file)

    result_dir_path = dir_path + "\\results_structure\\"

    if os.path.exists(result_dir_path):
        shutil.rmtree(result_dir_path)
    os.makedirs(result_dir_path)
    # file_path = result_dir_path + os.path.split(file)[1][0:-5]

    with open(result_dir_path + "structure.json", "w", encoding='utf-8') as write_file:
        json.dump(data, write_file, indent=4, ensure_ascii=False)
    with open(result_dir_path + "content.json", "w", encoding='utf-8') as write_file:
        json.dump(data1, write_file, indent=4, ensure_ascii=False)


'''''
def template(file):
    doc = read_docx(file)
    parsing = parse(doc)
    tem = open('template.txt').read()
    template = Template(tem)
    mydoc = Document()
    mydoc.add_paragraph(template.render(data=parsing))
    mydoc.save("result\\" + file[5:-5] + "_res.docx")
'''''

if __name__ == '__main__':
    folder_path = os.path.join(dir_path, "Files\\")
    if not os.path.exists(folder_path):
        sys.exit(1)

    for f in os.listdir(folder_path):
        # print(f)
        feed(os.path.join(folder_path, f))
